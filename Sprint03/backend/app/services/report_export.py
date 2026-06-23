from io import BytesIO

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

INK = colors.HexColor("#17362D")
ACCENT = colors.HexColor("#5D8B38")
LIGHT = colors.HexColor("#EEF3ED")
RED = colors.HexColor("#9B1C1C")


def _footer(canvas, document) -> None:
    canvas.saveState()
    canvas.setFont("Helvetica", 8)
    canvas.setFillColor(colors.HexColor("#66756F"))
    canvas.drawString(0.7 * inch, 0.45 * inch, "SecureGate AI - Master Security Report")
    canvas.drawRightString(7.8 * inch, 0.45 * inch, f"Page {document.page}")
    canvas.restoreState()


def generate_pdf(report: dict) -> bytes:
    buffer = BytesIO()
    document = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        rightMargin=0.7 * inch,
        leftMargin=0.7 * inch,
        topMargin=0.65 * inch,
        bottomMargin=0.65 * inch,
        title="SecureGate AI Deployment Approval",
        author="SecureGate AI",
    )
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="ReportTitle", parent=styles["Title"], fontName="Helvetica-Bold", fontSize=22, leading=26, textColor=INK, alignment=TA_CENTER, spaceAfter=8))
    styles.add(ParagraphStyle(name="Section", parent=styles["Heading2"], fontName="Helvetica-Bold", fontSize=13, leading=16, textColor=INK, spaceBefore=12, spaceAfter=7))
    styles.add(ParagraphStyle(name="Small", parent=styles["BodyText"], fontSize=8.5, leading=11, textColor=colors.HexColor("#50615B")))
    styles["BodyText"].fontSize = 9.5
    styles["BodyText"].leading = 13

    score = report["security_score"]["score"]
    decision = report["deployment_decision"]["decision"]
    summary = report["vulnerability_summary"]
    story = [
        Paragraph("SECUREGATE AI | COMPANY LOGO", styles["Small"]),
        Spacer(1, 0.08 * inch),
        Paragraph("Master Security Report", styles["ReportTitle"]),
        Paragraph(f"Scan ID: {report['executive_summary']['scan_id']}", styles["Small"]),
        Spacer(1, 0.16 * inch),
        Table(
            [["SECURITY SCORE", "DEPLOYMENT DECISION"], [f"{score}/100", decision]],
            colWidths=[3.25 * inch, 3.25 * inch],
            style=TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), LIGHT), ("TEXTCOLOR", (0, 0), (-1, 0), INK),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"), ("FONTSIZE", (0, 0), (-1, 0), 8),
                ("FONTNAME", (0, 1), (-1, 1), "Helvetica-Bold"), ("FONTSIZE", (0, 1), (-1, 1), 18),
                ("TEXTCOLOR", (1, 1), (1, 1), ACCENT if decision == "APPROVED" else RED),
                ("ALIGN", (0, 0), (-1, -1), "CENTER"), ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("BOX", (0, 0), (-1, -1), 0.75, colors.HexColor("#C9D4CE")),
                ("INNERGRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#DCE3DF")),
                ("TOPPADDING", (0, 0), (-1, -1), 8), ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
            ]),
        ),
        Paragraph("1. Executive Summary", styles["Section"]),
        Paragraph(report["executive_summary"]["summary"], styles["BodyText"]),
        Paragraph("2. Vulnerability Summary", styles["Section"]),
        Table(
            [["Critical", "High", "Medium", "Secrets", "Sonar Gate"], [summary.get("CRITICAL", 0), summary.get("HIGH", 0), summary.get("MEDIUM", 0), summary["secrets"], summary["sonarqube_quality_gate"]]],
            colWidths=[1.05 * inch, 1.05 * inch, 1.05 * inch, 1.05 * inch, 2.3 * inch],
            style=TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), INK), ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"), ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#CFD8D3")),
                ("TOPPADDING", (0, 0), (-1, -1), 7), ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
            ]),
        ),
        Paragraph("3. Deployment Decision", styles["Section"]),
    ]
    reasons = report["deployment_decision"]["reasons"] or ["All deployment approval conditions passed."]
    for reason in reasons:
        story.append(Paragraph(f"- {reason}", styles["BodyText"]))

    story.append(Paragraph("4. Tool Findings Summary", styles["Section"]))
    tool_rows = [["Tool", "Critical", "High", "Medium", "Low"]]
    for tool, counts in report["tool_findings_summary"].items():
        tool_rows.append([tool, counts.get("CRITICAL", 0), counts.get("HIGH", 0), counts.get("MEDIUM", 0), counts.get("LOW", 0)])
    if len(tool_rows) == 1:
        tool_rows.append(["No findings", 0, 0, 0, 0])
    story.append(Table(tool_rows, colWidths=[2.3 * inch, 1.05 * inch, 1.05 * inch, 1.05 * inch, 1.05 * inch], repeatRows=1, style=TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), INK), ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"), ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#CFD8D3")),
        ("ALIGN", (1, 1), (-1, -1), "CENTER"), ("TOPPADDING", (0, 0), (-1, -1), 6), ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
    ])))
    story.append(Paragraph("5. Recommendations", styles["Section"]))
    for item in report["recommendations"][:8]:
        actions = " ".join(item["actions"])
        story.append(Paragraph(f"<b>P{item['priority']} - {item['title']}:</b> {actions}", styles["BodyText"]))
    story.append(Paragraph("6. Report Metadata", styles["Section"]))
    metadata = report["report_metadata"]
    story.append(Paragraph(
        f"Project: {metadata['project_name']} | Branch: {metadata['branch']} | Commit: {metadata['commit_sha']}<br/>Generated: {metadata['generated_at']} | Report version: {metadata['report_version']}",
        styles["Small"],
    ))
    document.build(story, onFirstPage=_footer, onLaterPages=_footer)
    return buffer.getvalue()


def _set_cell_fill(cell, fill: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shading)


def generate_docx(report: dict) -> bytes:
    document = Document()
    section = document.sections[0]
    section.page_width, section.page_height = Inches(8.5), Inches(11)
    section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Inches(1)
    section.header_distance = section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name, normal.font.size, normal.font.color.rgb = "Calibri", Pt(10.5), RGBColor(25, 39, 34)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.05
    for style_name, size, color in [("Heading 1", 15, "2E7453"), ("Heading 2", 13, "2E7453"), ("Heading 3", 12, "214E3E")]:
        style = document.styles[style_name]
        style.font.name, style.font.size, style.font.bold = "Calibri", Pt(size), True
        style.font.color.rgb = RGBColor.from_string(color)

    header = section.header.paragraphs[0]
    header.text = "SecureGate AI | Master Security Report"
    header.runs[0].font.size = Pt(8)
    header.runs[0].font.color.rgb = RGBColor(95, 112, 105)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run("Confidential security assessment")
    footer.runs[0].font.size = Pt(8)

    logo = document.add_paragraph()
    logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_run = logo.add_run("SECUREGATE AI | COMPANY LOGO")
    logo_run.bold, logo_run.font.size, logo_run.font.color.rgb = True, Pt(9), RGBColor(95, 112, 105)
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(4)
    title_run = title.add_run("Master Security Report")
    title_run.bold, title_run.font.size, title_run.font.color.rgb = True, Pt(22), RGBColor(23, 54, 45)
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(f"Scan ID: {report['executive_summary']['scan_id']}")

    score_table = document.add_table(rows=2, cols=2)
    score_table.autofit = False
    for row in score_table.rows:
        row.cells[0].width = row.cells[1].width = Inches(3.25)
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    score_table.cell(0, 0).text, score_table.cell(0, 1).text = "SECURITY SCORE", "DEPLOYMENT DECISION"
    score_table.cell(1, 0).text = f"{report['security_score']['score']}/100"
    score_table.cell(1, 1).text = report["deployment_decision"]["decision"]
    for cell in score_table.rows[0].cells:
        _set_cell_fill(cell, "EEF3ED")
        cell.paragraphs[0].runs[0].bold = True
    for cell in score_table.rows[1].cells:
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(17)

    document.add_heading("1 Executive Summary", level=1)
    document.add_paragraph(report["executive_summary"]["summary"])
    document.add_heading("2 Security Score and Deployment Decision", level=1)
    reasons = report["deployment_decision"]["reasons"] or ["All deployment approval conditions passed."]
    document.add_paragraph(" ".join(reasons))
    document.add_heading("3 Vulnerability Summary", level=1)
    summary = report["vulnerability_summary"]
    document.add_paragraph(
        f"Critical: {summary.get('CRITICAL', 0)} | High: {summary.get('HIGH', 0)} | Medium: {summary.get('MEDIUM', 0)} | Secrets: {summary['secrets']} | SonarQube gate: {summary['sonarqube_quality_gate']}"
    )
    document.add_heading("4 Tool Findings Summary", level=1)
    table = document.add_table(rows=1, cols=5)
    table.autofit = False
    widths = [Inches(2.3), Inches(1.05), Inches(1.05), Inches(1.05), Inches(1.05)]
    for index, label in enumerate(["Tool", "Critical", "High", "Medium", "Low"]):
        table.cell(0, index).text = label
        table.cell(0, index).width = widths[index]
        _set_cell_fill(table.cell(0, index), "17362D")
        table.cell(0, index).paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        table.cell(0, index).paragraphs[0].runs[0].bold = True
    for tool, counts in report["tool_findings_summary"].items():
        cells = table.add_row().cells
        values = [tool, counts.get("CRITICAL", 0), counts.get("HIGH", 0), counts.get("MEDIUM", 0), counts.get("LOW", 0)]
        for index, value in enumerate(values):
            cells[index].text, cells[index].width = str(value), widths[index]
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cells[index].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT if index == 0 else WD_ALIGN_PARAGRAPH.CENTER
    document.add_heading("5 Recommendations", level=1)
    for item in report["recommendations"][:8]:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.add_run(f"P{item['priority']} - {item['title']}: ").bold = True
        paragraph.add_run(" ".join(item["actions"]))
    document.add_heading("6 Report Metadata", level=1)
    metadata = report["report_metadata"]
    document.add_paragraph(
        f"Project: {metadata['project_name']} | Branch: {metadata['branch']} | Commit: {metadata['commit_sha']}\nRepository: {metadata['repository_url']}\nGenerated: {metadata['generated_at']} | Version: {metadata['report_version']}"
    )
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()
